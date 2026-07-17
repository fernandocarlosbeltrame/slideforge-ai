from pathlib import Path
from uuid import uuid4
from zipfile import ZipFile
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image
from slideforge.domain.entities.content_block import ContentBlock
from slideforge.domain.entities.document_section import DocumentSection
from slideforge.domain.entities.source_document import SourceDocument
from slideforge.domain.enums.block_type import BlockType
from slideforge.domain.value_objects.table_data import TableCell, TableData, TableRow


class DocxReader:
    extensions = {".docx"}

    def supports(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def read(self, path: Path) -> SourceDocument:
        doc = Document(path)
        document = SourceDocument(id=str(uuid4()), name=path.name, source_path=str(path), format="docx", metadata=dict(doc.core_properties.__dict__))
        current = DocumentSection(id=str(uuid4()), title=path.stem, order=1)
        document.sections.append(current)
        media = self._extract_media(path)
        order = 1
        for element in doc.element.body.iterchildren():
            tag = element.tag
            if tag == qn("w:p"):
                paragraph = Paragraph(element, doc)
                text = (paragraph.text or "").strip()
                images = self._images_from_paragraph(paragraph, media)
                if text:
                    style = (paragraph.style.name or "").lower() if paragraph.style else ""
                    block_type, level = self._classify_paragraph(style, text, len(document.sections))
                    if block_type in {BlockType.TITLE, BlockType.SUBTITLE} and current.blocks:
                        current = DocumentSection(id=str(uuid4()), title=text, level=level or 1, order=len(document.sections) + 1)
                        document.sections.append(current)
                    elif block_type in {BlockType.TITLE, BlockType.SUBTITLE}:
                        current.title = text
                    current.blocks.append(ContentBlock(id=str(uuid4()), type=block_type, text=text.lstrip("•-* ").strip(), hierarchy_level=level, order=order, source_reference=f"paragraph:{order}", metadata={"style": style}))
                    order += 1
                for image in images:
                    image["order"] = order
                    current.blocks.append(ContentBlock(id=str(uuid4()), type=BlockType.IMAGE, text=image.get("alt_text") or f"Imagem {order}", order=order, source_reference=f"image:{order}", metadata=image))
                    order += 1
            elif tag == qn("w:tbl"):
                table = Table(element, doc)
                table_data = self._table_data(table)
                current.blocks.append(ContentBlock(id=str(uuid4()), type=BlockType.TABLE, text=f"Tabela {order}", order=order, source_reference=f"table:{order}", metadata={"table": table_data, "rows": table_data.as_matrix()}))
                order += 1
                for image in self._images_from_table(table, media):
                    image["order"] = order
                    current.blocks.append(ContentBlock(id=str(uuid4()), type=BlockType.IMAGE, text=image.get("alt_text") or f"Imagem {order}", order=order, source_reference=f"table-image:{order}", metadata=image))
                    order += 1
        return document

    @staticmethod
    def _classify_paragraph(style: str, text: str, section_count: int) -> tuple[BlockType, int]:
        if "title" in style or "título" in style:
            return (BlockType.TITLE if section_count == 1 else BlockType.SUBTITLE, 1)
        if "heading" in style or "cabeçalho" in style or "titulo" in style:
            level = next((int(token) for token in style.split() if token.isdigit()), 1)
            return (BlockType.SUBTITLE, level)
        if style.startswith("list") or text.startswith(("•", "-", "*")):
            return (BlockType.LIST_ITEM, 0)
        return (BlockType.PARAGRAPH, 0)

    @staticmethod
    def _extract_media(path: Path) -> dict[str, dict]:
        output_dir = path.parent / ".slideforge_assets" / path.stem
        output_dir.mkdir(parents=True, exist_ok=True)
        media: dict[str, dict] = {}
        try:
            with ZipFile(path) as archive:
                rels = {}
                rel_path = "word/_rels/document.xml.rels"
                if rel_path in archive.namelist():
                    import xml.etree.ElementTree as ET
                    root = ET.fromstring(archive.read(rel_path))
                    for rel in root:
                        rid = rel.attrib.get("Id")
                        target = rel.attrib.get("Target", "")
                        if rid and target.startswith("media/"):
                            rels[rid] = "word/" + target
                for rid, name in rels.items():
                    suffix = Path(name).suffix.lower()
                    if suffix not in {".png", ".jpg", ".jpeg", ".gif", ".bmp"}:
                        continue
                    target = output_dir / f"{rid}_{Path(name).name}"
                    target.write_bytes(archive.read(name))
                    width = height = 0
                    try:
                        with Image.open(target) as img:
                            width, height = img.size
                    except Exception:
                        pass
                    media[rid] = {"relationship_id": rid, "path": str(target), "format": suffix.lstrip("."), "width_px": width, "height_px": height, "aspect_ratio": width / height if height else None, "estimated_dpi": 96}
        except Exception:
            return {}
        return media

    @staticmethod
    def _images_from_paragraph(paragraph: Paragraph, media: dict[str, dict]) -> list[dict]:
        images: list[dict] = []
        for drawing in paragraph._element.xpath(".//w:drawing"):
            blips = drawing.xpath(".//a:blip")
            for blip in blips:
                rid = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
                if rid and rid in media:
                    item = dict(media[rid])
                    doc_pr = drawing.xpath(".//wp:docPr")
                    if doc_pr:
                        item["alt_text"] = doc_pr[0].get("descr") or doc_pr[0].get("title")
                    item["position"] = "inline"
                    images.append(item)
        return images

    def _images_from_table(self, table: Table, media: dict[str, dict]) -> list[dict]:
        images: list[dict] = []
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    images.extend(self._images_from_paragraph(paragraph, media))
        return images

    @staticmethod
    def _table_data(table: Table) -> TableData:
        rows: list[TableRow] = []
        for row in table.rows:
            rows.append(TableRow([TableCell(cell.text.strip()) for cell in row.cells]))
        has_header = bool(rows and any(cell.text for cell in rows[0].cells))
        return TableData(rows=rows, has_header=has_header)
