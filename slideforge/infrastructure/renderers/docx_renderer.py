from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from slideforge.application.publishing.speaker_notes import SpeakerNotesGenerator
from slideforge.domain.entities.content_audit import ContentAudit
from slideforge.domain.entities.presentation_plan import PresentationPlan
from slideforge.infrastructure.assets import AssetManager
from slideforge.infrastructure.renderers.serialization import plan_agenda
from slideforge.version import SLIDEFORGE_VERSION


class DocxRenderer:
    format_name = "docx"
    extension = ".docx"

    def render(self, plan: PresentationPlan, output_path: Path, *, audit: ContentAudit | None = None, assets: AssetManager | None = None, theme=None) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        _setup_styles(doc, theme)
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        _header_footer(section, plan.title)
        banner = assets.locate("banner") if assets else None
        if banner and Path(banner).exists():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(banner), width=Inches(5.8))
        title = doc.add_heading(plan.title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph("Publicação executiva editável")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"SlideForge {SLIDEFORGE_VERSION}").alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_page_break()
        doc.add_heading("Sumário", level=1)
        for seq, title_text in plan_agenda(plan):
            doc.add_paragraph(f"{seq}. {title_text}", style="List Bullet")
        notes = {n.slide_sequence: n for n in SpeakerNotesGenerator().generate(plan)}
        for slide in plan.slides:
            doc.add_page_break()
            doc.add_heading(f"{slide.sequence}. {slide.title}", level=1)
            if slide.subtitle:
                doc.add_paragraph(slide.subtitle, style="Intense Quote")
            for component in slide.components:
                _render_component(doc, component)
            note = notes[slide.sequence]
            doc.add_paragraph("Observações do slide", style="Heading 2")
            doc.add_paragraph(note.summary)
            doc.add_paragraph(f"Blocos utilizados: {', '.join(note.source_block_ids)}")
        if audit:
            doc.add_page_break()
            doc.add_heading("Auditoria técnica", level=1)
            for line in audit.as_report().splitlines():
                doc.add_paragraph(line)
        doc.save(output_path)
        return output_path


def _setup_styles(doc: Document, theme) -> None:
    primary = RGBColor(0, 74, 143)
    for style_name in ("Normal", "Body Text"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Aptos"
            style.font.size = Pt(10.5)
    for style_name in ("Heading 1", "Heading 2", "Title"):
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Aptos Display"
            style.font.color.rgb = primary


def _header_footer(section, title: str) -> None:
    header = section.header.paragraphs[0]
    header.text = title
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer = section.footer.paragraphs[0]
    footer.text = "SlideForge | Publicação executiva"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _render_component(doc: Document, component: dict) -> None:
    ctype = component.get("type")
    if ctype == "comparison":
        table = doc.add_table(rows=2, cols=2)
        table.style = "Table Grid"
        table.cell(0, 0).text = "Antes"
        table.cell(0, 1).text = "Depois"
        table.cell(1, 0).text = "\n".join(component.get("before", []))
        table.cell(1, 1).text = "\n".join(component.get("after", []))
        doc.add_paragraph()
    elif ctype == "image" and component.get("path") and Path(component["path"]).exists():
        doc.add_picture(component["path"], width=Inches(5.8))
        if component.get("caption"):
            doc.add_paragraph(component["caption"], style="Caption")
    elif ctype == "table":
        rows = component.get("rows", [])
        if rows:
            table = doc.add_table(rows=len(rows), cols=max(len(row) for row in rows))
            table.style = "Table Grid"
            for r, row in enumerate(rows):
                for c, cell in enumerate(row):
                    table.cell(r, c).text = str(cell)
            doc.add_paragraph()
    elif ctype == "timeline":
        table = doc.add_table(rows=max(len(component.get("items", [])), 1), cols=1)
        table.style = "Table Grid"
        for r, item in enumerate(component.get("items", [])):
            table.cell(r, 0).text = str(item)
    else:
        items = component.get("items", [])
        if len(items) <= 4 and items:
            table = doc.add_table(rows=(len(items) + 1) // 2, cols=2)
            table.style = "Table Grid"
            for idx, item in enumerate(items):
                table.cell(idx // 2, idx % 2).text = str(item)
        else:
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")
