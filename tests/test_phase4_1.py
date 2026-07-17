import json
from pathlib import Path
from zipfile import ZipFile

import fitz
from PIL import Image
from docx import Document

from slideforge.application.publishing.consistency_validator import PublishingConsistencyValidator
from slideforge.application.publishing.manifest import ManifestBuilder
from slideforge.application.publishing.package_builder import PublishingPackageBuilder
from slideforge.application.use_cases.publish_presentation import PublishPresentationUseCase
from slideforge.application.services.content_auditor import ContentAuditor
from slideforge.application.services.presentation_planner import PresentationPlanner
from slideforge.domain.entities.content_block import ContentBlock
from slideforge.domain.entities.document_section import DocumentSection
from slideforge.domain.entities.source_document import SourceDocument
from slideforge.domain.enums.block_type import BlockType
from slideforge.infrastructure.renderers import DocxRenderer, HtmlRenderer, JsonRenderer, MarkdownRenderer, PdfRenderer
from slideforge.version import SLIDEFORGE_VERSION


def _plan(tmp_path: Path):
    image = tmp_path / "img.png"
    Image.new("RGB", (24, 16), color=(0, 74, 143)).save(image)
    blocks = [
        ContentBlock(id="t", type=BlockType.TITLE, text="Publicação Executiva", hierarchy_level=1, order=1, source_reference="test"),
        ContentBlock(id="b1", type=BlockType.LIST_ITEM, text="Primeiro ponto", hierarchy_level=2, order=2, source_reference="test"),
        ContentBlock(id="b2", type=BlockType.LIST_ITEM, text="Segundo ponto", hierarchy_level=2, order=3, source_reference="test"),
    ]
    section = DocumentSection(id="s", title="Publicação Executiva", order=1, blocks=blocks)
    doc = SourceDocument(id="doc", name="teste.txt", source_path=str(tmp_path / "teste.txt"), format="txt", sections=[section])
    Path(doc.source_path).write_text("Publicação Executiva", encoding="utf-8")
    plan = PresentationPlanner().create_plan(doc)
    plan.slides[-1].components.append({"type": "image", "path": str(image), "caption": "Imagem de teste"})
    audit = ContentAuditor().audit(plan)
    return plan, audit


def test_pdf_renderer_uses_16_9_page(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    out = PdfRenderer().render(plan, tmp_path / "out.pdf", audit=audit)
    doc = fitz.open(out)
    page = doc[0].rect
    ratio = round(page.width / page.height, 2)
    assert ratio == 1.78
    assert doc.page_count >= len(plan.slides)


def test_html_renderer_has_navigation_accessibility_and_embedded_image(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    out = HtmlRenderer().render(plan, tmp_path / "out.html", audit=audit)
    html = out.read_text(encoding="utf-8")
    assert "<nav aria-label=" in html
    assert "Modo apresentação" in html
    assert "data:image/png;base64" in html
    assert "@media (max-width:900px)" in html


def test_docx_renderer_has_cover_summary_and_notes(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    out = DocxRenderer().render(plan, tmp_path / "out.docx", audit=audit)
    doc = Document(out)
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Sumário" in text
    assert "Observações do slide" in text
    assert "SlideForge" in text


def test_markdown_renderer_has_metadata_anchors_and_traceability_comments(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    out = MarkdownRenderer().render(plan, tmp_path / "out.md", audit=audit)
    md = out.read_text(encoding="utf-8")
    assert "slideforge_version" in md
    assert '<a id="slide-' in md
    assert "<!-- layout:" in md
    assert "<details>" in md


def test_manifest_and_package_are_created_with_hashes(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    docs = []
    for renderer in [JsonRenderer(), MarkdownRenderer(), HtmlRenderer()]:
        path = renderer.render(plan, tmp_path / f"deck{renderer.extension}", audit=audit)
        docs.append(type("Doc", (), {"format_name": renderer.format_name, "path": path})())
    consistency = PublishingConsistencyValidator().validate(plan, docs, audit)
    manifest = ManifestBuilder().build(plan, docs, tmp_path / "deck.manifest.json", audit=audit, consistency=consistency)
    data = json.loads(manifest.read_text(encoding="utf-8"))
    assert data["slideforge_version"] == SLIDEFORGE_VERSION
    assert all(item["sha256"] for item in data["formats"])
    package = PublishingPackageBuilder().build(tmp_path / "deck_package.zip", docs, manifest)
    with ZipFile(package) as zipf:
        names = set(zipf.namelist())
    assert "presentation/presentation.html" in names
    assert "presentation/presentation.manifest.json" in names


def test_consistency_validator_blocks_unused_content(tmp_path: Path):
    plan, audit = _plan(tmp_path)
    audit.unused_block_ids.add("x")
    out = JsonRenderer().render(plan, tmp_path / "out.json", audit=audit)
    docs = [type("Doc", (), {"format_name": "json", "path": out})()]
    result = PublishingConsistencyValidator().validate(plan, docs, audit)
    assert result.critical_count == 1


def test_rc1_public_outputs_do_not_contain_absolute_paths(tmp_path: Path):
    docx = tmp_path / "entrada.docx"
    doc = Document()
    doc.add_heading("Título", level=1)
    doc.add_paragraph("Texto")
    doc.save(docx)
    result = PublishPresentationUseCase(theme_name="corporate_blue").execute(docx, tmp_path / "saida", formats={"json", "markdown", "html"}, create_package=True)
    sensitive = ("C:\\Users", "C:/Users", "\\Users\\", "/Users/")
    for path in [d.path for d in result.documents] + [result.manifest_path]:
        text = Path(path).read_text(encoding="utf-8")
        assert not any(marker in text for marker in sensitive), path
    with ZipFile(result.package_path) as zipf:
        for name in zipf.namelist():
            assert not any(marker in name for marker in sensitive)
            if name.lower().endswith((".md", ".html", ".json", ".txt")):
                text = zipf.read(name).decode("utf-8")
                assert not any(marker in text for marker in sensitive), name




def test_output_stem_with_dotted_version_keeps_full_name(tmp_path: Path):
    from slideforge.application.publishing import output_with_extension

    output = output_with_extension(tmp_path / "slideforge_ai_demo_v1.0.0", ".pptx")
    assert output.name == "slideforge_ai_demo_v1.0.0.pptx"
