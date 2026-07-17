import io
import json
from urllib import error

import pytest

from slideforge.application.ai import AIConfig, AIContext, AISummarizationService, SummaryOptions, SummarizationEligibilityPolicy
from slideforge.domain.entities.content_block import ContentBlock
from slideforge.domain.enums.block_type import BlockType
from slideforge.infrastructure.ai import AIProviderFactory, FakeAIProvider, OllamaProvider
from slideforge.infrastructure.ai.ollama_provider import OllamaContentSummarizer, sanitize_external_error


class FakeHTTPResponse:
    def __init__(self, payload: str):
        self.payload = payload.encode("utf-8")

    def __enter__(self):
        return self

    def __exit__(self, exc_type, exc, tb):
        return False

    def read(self):
        return self.payload


def long_paragraph() -> str:
    return (
        "Em 2026, a empresa deverá revisar os procedimentos fiscais relacionados ao IBS, CBS e ISS, "
        "mantendo percentuais, datas, bases legais e valores íntegros durante qualquer análise. "
    ) * 12


def block(block_type: BlockType, text: str) -> ContentBlock:
    return ContentBlock(id="b1", type=block_type, text=text, hierarchy_level=1, order=1, source_reference="test")


def test_ai_config_disabled_by_default_and_supports_new_schema():
    config = AIConfig.from_dict({})
    assert config.enabled is False
    assert config.provider == "fake"
    assert config.active_provider == "fake"
    assert config.base_url == "http://localhost:11434"
    assert config.fallback_provider == "fake"


def test_provider_factory_resolves_fake_and_ollama():
    assert isinstance(AIProviderFactory().create(AIConfig(provider="fake")), FakeAIProvider)
    assert isinstance(AIProviderFactory().create(AIConfig(provider="ollama", model="llama3")), OllamaProvider)


def test_ollama_valid_response_is_parsed(monkeypatch):
    captured = {}

    def fake_urlopen(req, timeout):
        captured["timeout"] = timeout
        captured["body"] = json.loads(req.data.decode("utf-8"))
        return FakeHTTPResponse(json.dumps({"response": "Resumo seguro com 2026 e 12%."}))

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", fake_urlopen)
    config = AIConfig(enabled=True, provider="ollama", model="llama3", timeout_seconds=7, max_tokens=120)
    result = OllamaContentSummarizer(config).summarize_content(long_paragraph(), AIContext(purpose="teste"), options=SummaryOptions(max_length=120))
    assert result.success is True
    assert result.provider == "ollama"
    assert result.model == "llama3"
    assert result.text == "Resumo seguro com 2026 e 12%."
    assert captured["timeout"] == 7
    assert captured["body"]["stream"] is False
    assert "Preserve números" in captured["body"]["prompt"]
    assert "referências legais" in captured["body"]["prompt"]


def test_ollama_timeout_returns_failed_result(monkeypatch):
    def fake_urlopen(req, timeout):
        raise TimeoutError("timed out")

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", fake_urlopen)
    result = OllamaContentSummarizer(AIConfig(provider="ollama", model="llama3")).summarize_content("texto longo")
    assert result.success is False
    assert result.error_type == "timeout"
    assert result.text == "texto longo"


def test_ollama_connection_refused_returns_failed_result(monkeypatch):
    def fake_urlopen(req, timeout):
        raise error.URLError(ConnectionRefusedError("refused"))

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", fake_urlopen)
    result = OllamaContentSummarizer(AIConfig(provider="ollama", model="llama3")).summarize_content("texto")
    assert result.success is False
    assert result.error_type == "connection_unavailable"


def test_ollama_model_not_found_returns_failed_result(monkeypatch):
    def fake_urlopen(req, timeout):
        fp = io.BytesIO(b'{"error":"model not found"}')
        raise error.HTTPError(req.full_url, 404, "not found", hdrs=None, fp=fp)

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", fake_urlopen)
    result = OllamaContentSummarizer(AIConfig(provider="ollama", model="missing")).summarize_content("texto")
    assert result.success is False
    assert result.error_type == "model_not_found"


def test_ollama_empty_and_invalid_json_are_failed(monkeypatch):
    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", lambda req, timeout: FakeHTTPResponse("not-json"))
    invalid = OllamaContentSummarizer(AIConfig(provider="ollama", model="llama3")).summarize_content("texto")
    assert invalid.error_type == "invalid_response"

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", lambda req, timeout: FakeHTTPResponse(json.dumps({"response": "   "})))
    empty = OllamaContentSummarizer(AIConfig(provider="ollama", model="llama3")).summarize_content("texto")
    assert empty.error_type == "empty_response"


def test_summarization_service_uses_fallback_and_preserves_original(monkeypatch):
    def fake_urlopen(req, timeout):
        raise error.URLError(ConnectionRefusedError("refused"))

    monkeypatch.setattr("slideforge.infrastructure.ai.ollama_provider.request.urlopen", fake_urlopen)
    original = long_paragraph()
    config = AIConfig(enabled=True, provider="ollama", model="llama3", fallback_provider="fake")
    service = AISummarizationService(config, AIProviderFactory())
    result = service.summarize_block(block(BlockType.PARAGRAPH, original), AIContext(purpose="teste"))
    assert result.original_text == original
    assert result.summary.fallback_used is True
    assert result.summary.status == "fallback_fake"
    assert "fallback fake" in "; ".join(result.summary.warnings).lower()


def test_disabled_ai_keeps_current_behavior():
    original = long_paragraph()
    service = AISummarizationService(AIConfig(enabled=False, provider="ollama", model="llama3"), AIProviderFactory())
    result = service.summarize_block(block(BlockType.PARAGRAPH, original))
    assert result.derived_summary == original
    assert result.summary.status == "disabled"


def test_eligibility_policy_blocks_titles_tables_dates_and_short_lists():
    policy = SummarizationEligibilityPolicy()
    assert policy.evaluate_block(block(BlockType.TITLE, "Título grande qualquer")).eligible is False
    assert policy.evaluate_block(block(BlockType.TABLE, "A|B\n1|2")).eligible is False
    assert policy.evaluate_block(block(BlockType.PARAGRAPH, "12/05/2026")).eligible is False
    assert policy.evaluate_block(block(BlockType.LIST_ITEM, "Item curto")).eligible is False
    assert policy.evaluate_block(block(BlockType.PARAGRAPH, long_paragraph())).eligible is True


def test_sanitized_logs_do_not_expose_local_paths_or_text():
    message = r"Erro em C:\Users\Pessoa\Documents\arquivo.docx com conteúdo confidencial muito longo " + "x" * 300
    sanitized = sanitize_external_error(message)
    assert "Pessoa" not in sanitized
    assert "C:\\Users\\Pessoa" not in sanitized
    assert len(sanitized) <= 180



def test_publish_use_case_does_not_apply_ai_when_disabled(tmp_path):
    from docx import Document
    from slideforge.application.use_cases.publish_presentation import PublishPresentationUseCase

    docx = tmp_path / "entrada.docx"
    doc = Document()
    doc.add_heading("Título", level=1)
    doc.add_paragraph(long_paragraph())
    doc.save(docx)

    result = PublishPresentationUseCase(ai_config=AIConfig(enabled=False)).execute(docx, tmp_path / "saida", formats={"json"}, create_package=False)
    blocks = [block for section in result.plan.source_document.sections for block in section.blocks]
    assert all("ai_summary" not in block.metadata for block in blocks)


def test_publish_use_case_stores_ai_summary_as_derived_metadata(tmp_path):
    from docx import Document
    from slideforge.application.use_cases.publish_presentation import PublishPresentationUseCase

    original = long_paragraph()
    docx = tmp_path / "entrada.docx"
    doc = Document()
    doc.add_heading("Título", level=1)
    doc.add_paragraph(original)
    doc.save(docx)

    config = AIConfig(enabled=True, provider="fake", model="fake-deterministic-v1")
    result = PublishPresentationUseCase(ai_config=config, ai_provider_factory=AIProviderFactory()).execute(docx, tmp_path / "saida", formats={"json"}, create_package=False)
    paragraph_blocks = [block for section in result.plan.source_document.sections for block in section.blocks if block.type == BlockType.PARAGRAPH]
    assert paragraph_blocks
    assert paragraph_blocks[0].text == original.strip()
    assert paragraph_blocks[0].metadata["ai_summary"]["status"] == "fake_summary"
    assert paragraph_blocks[0].metadata["ai_summary"]["summary"] != original

